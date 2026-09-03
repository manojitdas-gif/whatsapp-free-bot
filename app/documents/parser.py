"""
parser.py — Unified attachment and document processing router.
"""

import os
from typing import Tuple
from app.documents.pdf_parser import extract_pdf_text
from app.documents.excel_parser import extract_spreadsheet_text
from app.documents.ocr_engine import run_image_ocr

def parse_attachment(file_path: str) -> Tuple[str, bool]:
    """
    Parses any supported file (PDF, Excel, Word, CSV, Images) and returns:
        (extracted_text: str, success: bool)
    """
    if not file_path or not os.path.exists(file_path):
        return "", False

    ext = os.path.splitext(file_path)[1].lower()

    if ext in (".png", ".jpg", ".jpeg", ".webp", ".bmp"):
        text = run_image_ocr(file_path)
        return text, bool(text)

    elif ext == ".pdf":
        text = extract_pdf_text(file_path)
        if not text:
            # Scanned PDF fallback
            text = run_image_ocr(file_path)
        return text, bool(text)

    elif ext in (".xlsx", ".xls", ".csv"):
        text = extract_spreadsheet_text(file_path)
        return text, bool(text)

    elif ext in (".docx", ".doc"):
        try:
            import docx
            doc = docx.Document(file_path)
            full_text = [p.text for p in doc.paragraphs if p.text]
            for table in doc.tables:
                for row in table.rows:
                    full_text.append(" | ".join(cell.text.strip() for cell in row.cells if cell.text.strip()))
            text = "\n".join(full_text)
            return text, bool(text)
        except Exception:
            return "", False

    elif ext == ".txt":
        try:
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                return f.read(), True
        except Exception:
            return "", False

    return "", False
