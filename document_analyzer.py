"""
document_analyzer.py — Universal Document, Photo, Screenshot & Text Analyzer.

SUPPORTED FORMATS:
  1. Photos & Screenshots: .jpg, .jpeg, .png, .webp, .bmp, .tiff, .jfif (via Windows Native OCR)
  2. PDF Documents: .pdf (all pages via pypdf)
  3. Excel Spreadsheets: .xlsx (openpyxl) and .xls (xlrd)
  4. CSV & TSV Files: .csv, .tsv, .txt
  5. Word Documents: .docx (python-docx)
  6. WhatsApp Chat Texts: Multiline catalogs, product lists, specifications, quantities

Extracts:
  - All product names, descriptions, and line items
  - Quantities & units (pcs, kg, meters, rolls, boxes, sets, nos, cartons, packets, etc.)
  - Dimensions & sizes (e.g. 12x10x8, 25 sq mm, 100mm, 2 inch, 5 ply, etc.)
  - Rates, prices, MRP, and discount terms
  - Business details (Company Name, GSTIN, Address, Contact details)
"""

import os
import re
import csv
import subprocess
try:
    import pypdf
except ImportError:
    pypdf = None

try:
    import openpyxl
except ImportError:
    openpyxl = None

try:
    import xlrd
except ImportError:
    xlrd = None

try:
    import docx
except ImportError:
    docx = None

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
OCR_SCRIPT = os.path.join(BASE_DIR, "ocr_engine.ps1")
DOWNLOADS_DIR = os.path.join(BASE_DIR, "data", "customer_files")
os.makedirs(DOWNLOADS_DIR, exist_ok=True)


# ─── 1. IMAGE & SCREENSHOT OCR ───────────────────────────────────────────────

def run_image_ocr(image_path: str) -> str:
    """Run native Windows OCR on an image/screenshot file and return extracted text."""
    if not os.path.exists(image_path):
        return ""

    cmd = [
        "powershell",
        "-NoProfile",
        "-ExecutionPolicy", "Bypass",
        "-File", OCR_SCRIPT,
        "-ImagePath", image_path,
    ]

    try:
        proc = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            timeout=30,
        )
        output = proc.stdout.strip()
        lines = [line.strip() for line in output.splitlines() if line.strip()]
        return "\n".join(lines)
    except Exception as e:
        print(f"[OCR ERROR] Image OCR failed: {e}")
        return ""


# ─── 2. PDF EXTRACTION ────────────────────────────────────────────────────────

def extract_pdf_text(pdf_path: str) -> str:
    """Extract text from all pages of a PDF document."""
    if not os.path.exists(pdf_path):
        return ""

    try:
        reader = pypdf.PdfReader(pdf_path)
        extracted = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            if text.strip():
                extracted.append(f"[Page {i + 1}]\n" + text.strip())
        return "\n\n".join(extracted).strip()
    except Exception as e:
        print(f"[PDF ERROR] PDF extraction failed: {e}")
        return ""


# ─── 3. EXCEL SPREADSHEETS (.xlsx & .xls) ────────────────────────────────────

def extract_excel_text(excel_path: str) -> str:
    """Extract tabular product lists from Excel (.xlsx and .xls) files."""
    if not os.path.exists(excel_path):
        return ""

    ext = os.path.splitext(excel_path)[1].lower()
    extracted_lines = []

    # Handle .xlsx with openpyxl
    if ext == ".xlsx":
        try:
            wb = openpyxl.load_workbook(excel_path, data_only=True)
            for sheet in wb.sheetnames:
                ws = wb[sheet]
                extracted_lines.append(f"--- Sheet: {sheet} ---")
                for row in ws.iter_rows(values_only=True):
                    # Filter out empty cells
                    cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
                    if cells:
                        extracted_lines.append(" | ".join(cells))
            return "\n".join(extracted_lines)
        except Exception as e:
            print(f"[EXCEL ERROR] .xlsx extraction failed: {e}")

    # Handle .xls with xlrd
    elif ext == ".xls":
        try:
            wb = xlrd.open_workbook(excel_path)
            for sheet in wb.sheets():
                extracted_lines.append(f"--- Sheet: {sheet.name} ---")
                for rx in range(sheet.nrows):
                    row_vals = [str(val).strip() for val in sheet.row_values(rx) if str(val).strip()]
                    if row_vals:
                        extracted_lines.append(" | ".join(row_vals))
            return "\n".join(extracted_lines)
        except Exception as e:
            print(f"[EXCEL ERROR] .xls extraction failed: {e}")

    return ""


# ─── 4. CSV & TSV FILES ───────────────────────────────────────────────────────

def extract_csv_text(csv_path: str) -> str:
    """Extract rows from CSV / TSV text files."""
    if not os.path.exists(csv_path):
        return ""

    try:
        extracted = []
        with open(csv_path, "r", encoding="utf-8-sig", errors="replace") as f:
            # Detect delimiter
            sample = f.read(2048)
            f.seek(0)
            delim = "\t" if "\t" in sample and "," not in sample else ","
            reader = csv.reader(f, delimiter=delim)
            for row in reader:
                cells = [c.strip() for c in row if c.strip()]
                if cells:
                    extracted.append(" | ".join(cells))
        return "\n".join(extracted)
    except Exception as e:
        print(f"[CSV ERROR] CSV extraction failed: {e}")
        return ""


# ─── 5. WORD DOCUMENTS (.docx) ────────────────────────────────────────────────

def extract_docx_text(docx_path: str) -> str:
    """Extract all text and table contents from Word .docx documents."""
    if not os.path.exists(docx_path):
        return ""

    try:
        doc = docx.Document(docx_path)
        extracted = []

        # Paragraphs
        for p in doc.paragraphs:
            if p.text.strip():
                extracted.append(p.text.strip())

        # Tables
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    # Remove consecutive duplicates in merged cells
                    unique_cells = []
                    for c in cells:
                        if not unique_cells or unique_cells[-1] != c:
                            unique_cells.append(c)
                    extracted.append(" | ".join(unique_cells))

        return "\n".join(extracted)
    except Exception as e:
        print(f"[DOCX ERROR] Word extraction failed: {e}")
        return ""


# ─── 6. SMART DETAILED REQUIREMENTS PARSER ───────────────────────────────────

def parse_product_details(raw_text: str) -> str:
    """
    Intelligently analyzes raw text from any source (Text, OCR, PDF, Excel, Word).
    Extracts all products, quantities, dimensions, units, rates, and specifications.
    Returns a clean, structured multi-line summary of all detected requirements.
    """
    if not raw_text or not raw_text.strip():
        return ""

    lines = [line.strip() for line in raw_text.splitlines() if line.strip()]
    if not lines:
        return ""

    product_items = []
    quantities = []
    specifications = []

    qty_pattern = re.compile(
        r'(\d+(?:\.\d+)?[\s]*(?:pcs|pc|nos|pieces|boxes|box|kg|g|gm|ton|tonne|m|meter|mtr|cm|inch|in|ft|feet|set|sets|units|packet|pkt|pkts|roll|rolls|carton|cartons|dozen|doz|bundle|bundles|strip|strips|sq\s*mm|sqmm|sq\s*ft|amp|kva|watt|w|v|volt))',
        re.IGNORECASE
    )
    dim_pattern = re.compile(
        r'(\d+[\s]*[xX*×][\s]*\d+[\s]*(?:[xX*×][\s]*\d+)?[\s]*(?:mm|cm|inch|in|ft|m)?)',
        re.IGNORECASE
    )
    rate_pattern = re.compile(
        r'(?:rs|inr|rate|price|mrp|\u20b9)[\s.:]*(\d+(?:\.\d{1,2})?)',
        re.IGNORECASE
    )

    req_keywords = [
        "product", "item", "mcb", "socket", "box", "carton", "cable", "lugs", "wire",
        "roll", "tape", "switch", "model", "enclosure", "heavy duty", "corrugated",
        "ply", "copper", "aluminium", "brass", "steel", "plastic", "paper", "packaging",
        "dimension", "quantity", "specification", "catalogue", "catalog"
    ]

    exclude_keywords = [
        "--- page", "[page", "business info", "gst", "gstin", "email", "www.", "http",
        "document content", "thanks", "thank you", "contact person", "dear sir", "hello",
        "regards", "since :", "pin:", "road", "street", "bazar", "bazaar", "lane",
        "where is", "call me", "how much", "send quote", "send quotation"
    ]

    for line in lines:
        line_lower = line.lower()

        # Skip headers / page markers / separator lines / non-product lines
        if line.startswith("---") or line.startswith("[Page") or len(line) < 2:
            continue
        if any(ex in line_lower for ex in exclude_keywords):
            continue
        if line_lower in ("hi", "hello", "hey", "good morning", "please send", "ok", "ok thanks", "thanks", "yes", "no"):
            continue

        has_qty = bool(qty_pattern.search(line))
        has_dim = bool(dim_pattern.search(line))
        has_rate = bool(rate_pattern.search(line))
        has_kw = any(kw in line_lower for kw in req_keywords)

        # Must have at least one product indicator: unit, dimension, rate, or product keyword
        if has_qty or has_dim or has_rate or has_kw:
            clean_item = re.sub(r'^(?:\d+[\.\)]|[•\-\*\+])\s+', '', line).strip()
            if clean_item and clean_item not in product_items:
                product_items.append(clean_item)

    # Format pure product items only
    if product_items:
        return "\n".join(f"• {item}" for item in product_items[:12])

    return ""


# ─── 7. UNIVERSAL FILE ANALYZER ENTRY POINT ───────────────────────────────────

def analyze_file(file_path: str, file_type: str = "") -> tuple[str, str]:
    """
    Universal entry point to analyze ANY customer file:
    - Photos / Screenshots (.jpg, .jpeg, .png, .webp, .bmp, .tiff, .jfif)
    - PDFs (.pdf)
    - Excel Spreadsheets (.xlsx, .xls)
    - CSV / TSV text (.csv, .tsv, .txt)
    - Word Documents (.docx)

    Returns:
        (structured_product_summary, full_raw_text)
    """
    if not os.path.exists(file_path):
        return ("File not found on disk", "")

    ext = os.path.splitext(file_path)[1].lower()
    raw_text = ""

    print(f"[ANALYZER] Processing file '{os.path.basename(file_path)}' (Format: {ext or file_type})...")

    # 1. Images & Screenshots
    if ext in (".jpg", ".jpeg", ".png", ".webp", ".bmp", ".tiff", ".jfif") or "image" in file_type.lower():
        raw_text = run_image_ocr(file_path)

    # 2. PDF Documents
    elif ext == ".pdf" or "pdf" in file_type.lower():
        raw_text = extract_pdf_text(file_path)

    # 3. Excel Spreadsheets (.xlsx, .xls)
    elif ext in (".xlsx", ".xls") or "sheet" in file_type.lower() or "excel" in file_type.lower():
        raw_text = extract_excel_text(file_path)

    # 4. CSV & TSV
    elif ext in (".csv", ".tsv"):
        raw_text = extract_csv_text(file_path)

    # 5. Word Documents (.docx)
    elif ext in (".docx", ".doc") or "word" in file_type.lower():
        raw_text = extract_docx_text(file_path)

    # 6. Plain Text Files (.txt)
    elif ext in (".txt", ".text"):
        try:
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                raw_text = f.read()
        except Exception:
            pass

    # 7. Fallback: try image OCR if unrecognized binary
    else:
        raw_text = run_image_ocr(file_path)

    summary = parse_product_details(raw_text)
    return summary, raw_text
